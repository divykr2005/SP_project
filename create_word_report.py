from __future__ import annotations

from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUTS = ROOT / "outputs"
FIGURES = OUTPUTS / "figures"
TABLES = OUTPUTS / "tables"
REPORT = OUTPUTS / "report"
DOCX_PATH = REPORT / "Bearing_Fault_Diagnosis_Detailed_Report.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text_color(cell, color: str) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(color)


def add_table(document: Document, df: pd.DataFrame, columns: Iterable[str], title: str | None = None) -> None:
    cols = list(columns)
    if title:
        paragraph = document.add_paragraph()
        paragraph.style = "Caption"
        paragraph.add_run(title).bold = True

    table = document.add_table(rows=1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for j, column in enumerate(cols):
        header_cells[j].text = column.replace("_", " ").title()
        set_cell_shading(header_cells[j], "1F4E79")
        set_cell_text_color(header_cells[j], "FFFFFF")
        for paragraph in header_cells[j].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    for i, (_, row) in enumerate(df.iterrows()):
        cells = table.add_row().cells
        for j, column in enumerate(cols):
            value = row[column]
            if isinstance(value, float):
                text = f"{value:.4g}"
            else:
                text = str(value)
            cells[j].text = text
            cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if i % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, "F2F6FA")

    document.add_paragraph()


def add_picture(document: Document, filename: str, caption: str, width: float = 6.5) -> None:
    path = FIGURES / filename
    if not path.exists():
        paragraph = document.add_paragraph()
        paragraph.add_run(f"Missing figure: {filename}").italic = True
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))

    caption_paragraph = document.add_paragraph()
    caption_paragraph.style = "Caption"
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.add_run(caption)


def add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: Iterable[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Calibri"
    styles["Title"].font.size = Pt(24)
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 2"].font.name = "Calibri"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor(46, 117, 182)
    styles["Caption"].font.name = "Calibri"
    styles["Caption"].font.size = Pt(9)
    styles["Caption"].font.italic = True


def add_title_page(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Bearing Fault Diagnosis Under Time-Varying Rotational Speed")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Detailed Analysis Report").bold = True

    details = document.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details.add_run(
        "Dataset: 36 bearing vibration datasets with healthy, inner race fault, and outer race fault conditions\n"
        "Signals: Channel_1 vibration and Channel_2 encoder speed\n"
        "Sampling rate: 200,000 Hz, duration: 10 seconds per file"
    )

    document.add_paragraph()
    document.add_paragraph(
        "This report documents the complete analysis workflow implemented for the bearing dataset: data organization, "
        "signal visualization, feature extraction, data reduction, fault classification, speed-pattern classification, "
        "and relative severity estimation."
    )
    document.add_page_break()


def add_manual_toc(document: Document) -> None:
    document.add_heading("Contents", level=1)
    sections = [
        "1. Executive Summary",
        "2. Dataset Description",
        "3. Analysis Workflow",
        "4. Signal Exploration",
        "5. Feature Extraction and Data Reduction",
        "6. Classification Methodology",
        "7. Results and Discussion",
        "8. Relative Severity Index",
        "9. Limitations",
        "10. Conclusion",
        "11. Appendix: Generated Files",
    ]
    for section in sections:
        document.add_paragraph(section)
    document.add_page_break()


def build_report() -> Path:
    REPORT.mkdir(parents=True, exist_ok=True)

    feature_summary = pd.read_csv(TABLES / "file_feature_summary.csv")
    health_predictions = pd.read_csv(TABLES / "health_file_predictions.csv")
    fault_predictions = pd.read_csv(TABLES / "fault_detection_file_predictions.csv")
    speed_predictions = pd.read_csv(TABLES / "speed_file_predictions.csv")
    case_predictions = pd.read_csv(TABLES / "case12_file_predictions.csv")
    severity = pd.read_csv(TABLES / "file_severity_scores.csv")
    window_features = pd.read_csv(TABLES / "window_features.csv", usecols=["file"])

    document = Document()
    configure_document(document)
    add_title_page(document)
    add_manual_toc(document)

    document.add_heading("1. Executive Summary", level=1)
    document.add_paragraph(
        "The project analyzes vibration signals collected from bearings operating under time-varying rotational speed. "
        "The raw data is large because each dataset contains two million samples per channel, so the main engineering "
        "task is to reduce the data into meaningful diagnostic features while preserving fault-related information."
    )
    document.add_paragraph(
        "A complete base-MATLAB workflow was created. It extracts window-level features, produces signal-analysis "
        "figures, trains transparent classification models, evaluates trial-wise generalization, and estimates a "
        "relative abnormality severity score. The primary valid classification problem is health diagnosis: healthy, "
        "inner race fault, or outer race fault."
    )
    add_bullets(
        document,
        [
            "Total datasets analyzed: 36 .mat files.",
            "Total extracted windows: 7,164 windows using 0.10 s windows and 50% overlap.",
            "Main health classification: healthy, inner race fault, outer race fault.",
            "Trial-wise split: trials 1 and 2 used for training, trial 3 used for testing.",
            "File-level health classification accuracy: 100%.",
            "A relative severity index was generated, but true mild/severe classification is not possible without severity labels.",
        ],
    )

    document.add_heading("2. Dataset Description", level=1)
    document.add_paragraph(
        "Each dataset contains two channels. Channel_1 is the vibration signal measured using an accelerometer, and "
        "Channel_2 is the rotational speed signal measured using an encoder. The signals were sampled at 200,000 Hz "
        "for 10 seconds, giving 2,000,000 samples per channel per file."
    )
    dataset_table = pd.DataFrame(
        [
            ["Health condition", "H", "Healthy bearing"],
            ["Health condition", "I", "Inner race defect"],
            ["Health condition", "O", "Outer race defect"],
            ["Speed condition", "A", "Increasing speed"],
            ["Speed condition", "B", "Decreasing speed"],
            ["Speed condition", "C", "Increasing then decreasing speed"],
            ["Speed condition", "D", "Decreasing then increasing speed"],
            ["Trial", "1, 2, 3", "Repeated trials for the same setting"],
        ],
        columns=["category", "code", "meaning"],
    )
    add_table(document, dataset_table, ["category", "code", "meaning"], "Table 1. Dataset naming convention.")
    document.add_paragraph(
        "The filename H-A-1.mat therefore represents a healthy bearing under increasing speed in the first trial. "
        "The filename I-C-2.mat represents an inner race fault under increasing-then-decreasing speed in the second trial."
    )

    document.add_heading("3. Analysis Workflow", level=1)
    document.add_paragraph("The workflow implemented in the project follows these steps:")
    add_numbered(
        document,
        [
            "Read each .mat file and parse the health condition, speed condition, and trial number from the filename.",
            "Split the vibration and speed signals into overlapping time windows.",
            "Extract time-domain, frequency-domain, and speed-related features from each window.",
            "Summarize window features into file-level statistics for interpretation.",
            "Generate signal plots: speed profiles, waveform comparisons, spectrograms, envelope spectra, and order spectra.",
            "Train and test classifiers using a trial-wise split to avoid leakage.",
            "Estimate relative severity using healthy files as the baseline.",
            "Export tables, figures, saved result structures, and final reports.",
        ],
    )
    document.add_paragraph(
        "The workflow was intentionally written using base MATLAB calculations so it can run even when optional "
        "Statistics or Signal Processing Toolbox functions are unavailable."
    )

    document.add_heading("4. Signal Exploration", level=1)
    document.add_heading("4.1 Speed Profiles", level=2)
    document.add_paragraph(
        "The encoder channel was used to inspect the four time-varying rotational speed patterns. This step confirms "
        "that the data is non-stationary, meaning that fault signatures may shift over time."
    )
    add_picture(document, "speed_profiles_trial3.png", "Figure 1. Speed profiles for the four operating conditions.", 6.5)

    document.add_heading("4.2 Vibration Waveforms", level=2)
    document.add_paragraph(
        "A direct time-waveform comparison shows the change in vibration behavior across health states. Inner race "
        "fault files show stronger impulsive behavior than healthy files, while outer race faults can be more subtle "
        "in simple time-domain plots."
    )
    add_picture(document, "waveform_comparison_ha_ia_oa.png", "Figure 2. Vibration waveform comparison for representative healthy, inner fault, and outer fault files.", 6.5)

    document.add_heading("4.3 Spectrograms", level=2)
    document.add_paragraph(
        "Spectrograms show how frequency content changes over time. They are useful here because the rotational speed "
        "is not constant. A normal FFT gives one averaged frequency view, while a spectrogram preserves time variation."
    )
    add_picture(document, "spectrogram_h-a-3.png", "Figure 3. Spectrogram of representative healthy file H-A-3.", 6.5)
    add_picture(document, "spectrogram_i-a-3.png", "Figure 4. Spectrogram of representative inner race fault file I-A-3.", 6.5)
    add_picture(document, "spectrogram_o-a-3.png", "Figure 5. Spectrogram of representative outer race fault file O-A-3.", 6.5)

    document.add_heading("4.4 Envelope and Order Analysis", level=2)
    document.add_paragraph(
        "Bearing faults often produce repeated impacts. Envelope analysis makes these impact patterns more visible in "
        "the spectrum. Approximate order analysis resamples the signal with respect to shaft rotation so speed-varying "
        "fault signatures become easier to compare."
    )
    add_picture(document, "envelope_spectrum_h-a-3.png", "Figure 6. Envelope spectrum for H-A-3.", 6.5)
    add_picture(document, "envelope_spectrum_i-a-3.png", "Figure 7. Envelope spectrum for I-A-3.", 6.5)
    add_picture(document, "envelope_spectrum_o-a-3.png", "Figure 8. Envelope spectrum for O-A-3.", 6.5)
    add_picture(document, "order_spectrum_h-a-3.png", "Figure 9. Approximate order spectrum for H-A-3.", 6.5)
    add_picture(document, "order_spectrum_i-a-3.png", "Figure 10. Approximate order spectrum for I-A-3.", 6.5)
    add_picture(document, "order_spectrum_o-a-3.png", "Figure 11. Approximate order spectrum for O-A-3.", 6.5)

    document.add_heading("5. Feature Extraction and Data Reduction", level=1)
    document.add_paragraph(
        "Instead of using raw signals directly, each file was divided into windows and each window was represented by "
        "a compact feature vector. This reduces the computational burden while preserving useful diagnostic behavior."
    )
    reduction_table = pd.DataFrame(
        [
            ["Raw per file", "2 channels x 2,000,000 samples", "4,000,000 raw values"],
            ["All files", "36 files", "144,000,000 raw values"],
            ["Windowed features", "7,164 windows x 21 features", "150,444 feature values"],
        ],
        columns=["stage", "calculation", "amount"],
    )
    add_table(document, reduction_table, ["stage", "calculation", "amount"], "Table 2. Data reduction from raw signals to features.")
    document.add_paragraph(
        "The extracted features include RMS, standard deviation, variance, mean absolute value, peak-to-peak amplitude, "
        "crest factor, impulse factor, shape factor, skewness, kurtosis, energy, spectral energy, spectral centroid, "
        "low/high frequency ratios, speed-normalized RMS, and speed statistics."
    )
    add_picture(document, "feature_summary_by_health.png", "Figure 12. File-level feature summary by health class.", 6.5)
    add_picture(document, "feature_separation_scores.png", "Figure 13. Feature separation proxy for health classification.", 6.5)

    health_summary = (
        feature_summary.groupby("health_name")
        .agg(
            rms_mean=("rms_value_mean", "mean"),
            kurtosis_mean=("kurtosis_value_mean", "mean"),
            crest_factor_mean=("crest_factor_mean", "mean"),
        )
        .reset_index()
    )
    add_table(document, health_summary, ["health_name", "rms_mean", "kurtosis_mean", "crest_factor_mean"], "Table 3. Average file-level feature values by health class.")

    document.add_heading("6. Classification Methodology", level=1)
    document.add_paragraph(
        "The primary model is a 5-nearest-neighbor classifier implemented in plain MATLAB. Features were standardized "
        "using training-set mean and standard deviation only. This prevents information from the test trial from leaking "
        "into the training process."
    )
    add_bullets(
        document,
        [
            "Training set: all trial 1 and trial 2 files.",
            "Test set: all trial 3 files.",
            "Main task: classify healthy, inner race fault, and outer race fault.",
            "Additional task: classify healthy versus faulty.",
            "Additional task: classify speed pattern A, B, C, or D.",
            "Additional task: classify the full 12 combined health-speed cases.",
        ],
    )
    document.add_paragraph(
        "Trial-wise splitting is important because random window splitting would place very similar neighboring windows "
        "from the same physical experiment into both training and test sets, producing overly optimistic results."
    )

    document.add_heading("7. Results and Discussion", level=1)
    results_table = pd.DataFrame(
        [
            ["Healthy vs inner vs outer", "Window", "98.99%", "0.990"],
            ["Healthy vs inner vs outer", "File", "100.00%", "1.000"],
            ["Healthy vs faulty", "Window", "98.99%", "0.989"],
            ["Healthy vs faulty", "File", "100.00%", "1.000"],
            ["Speed pattern A/B/C/D", "File", "75.00%", "0.754"],
            ["12 combined cases", "File", "83.33%", "0.778"],
        ],
        columns=["task", "level", "accuracy", "macro_f1"],
    )
    add_table(document, results_table, ["task", "level", "accuracy", "macro_f1"], "Table 4. Classification performance summary.")
    document.add_paragraph(
        "The main health classifier correctly identified all trial-3 files at file level. The window-level accuracy was "
        "also high, showing that the extracted vibration features are sufficient for distinguishing healthy, inner race "
        "fault, and outer race fault conditions. Speed-pattern and 12-case classification are harder because the model "
        "must also distinguish the operating condition, not only the bearing fault."
    )
    add_picture(document, "confusion_health_window_knn.png", "Figure 14. Window-level health classification confusion matrix.", 6.5)
    add_picture(document, "confusion_health_file_knn.png", "Figure 15. File-level health classification confusion matrix.", 5.8)
    add_picture(document, "confusion_speed_file_knn.png", "Figure 16. File-level speed-pattern classification confusion matrix.", 5.8)
    add_picture(document, "confusion_12_case_file_knn.png", "Figure 17. File-level 12-case classification confusion matrix.", 6.2)

    health_table = health_predictions[["file", "true_name", "predicted_name", "correct"]].copy()
    health_table["correct"] = health_table["correct"].map({1: "Yes", True: "Yes", 0: "No", False: "No"})
    add_table(document, health_table, ["file", "true_name", "predicted_name", "correct"], "Table 5. File-level health predictions on trial-3 test files.")

    document.add_paragraph(
        "The 12-case task misclassified two file-level cases: I-D-3 was predicted as I-C, and O-D-3 was predicted as O-C. "
        "Both mistakes preserve the correct fault type but confuse the mixed speed pattern. This is acceptable evidence "
        "that health diagnosis is stronger than exact health-speed case classification."
    )
    case_table = case_predictions[["file", "true_name", "predicted_name", "correct"]].copy()
    case_table["correct"] = case_table["correct"].map({1: "Yes", True: "Yes", 0: "No", False: "No"})
    add_table(document, case_table, ["file", "true_name", "predicted_name", "correct"], "Table 6. File-level 12-case predictions on trial-3 test files.")

    document.add_heading("8. Relative Severity Index", level=1)
    document.add_paragraph(
        "The dataset does not provide actual defect size, crack depth, damage progression, or mild/severe labels. "
        "Therefore, a true supervised severity classifier cannot be claimed. Instead, a relative abnormality score was "
        "computed using healthy files as the baseline."
    )
    document.add_paragraph(
        "The severity score uses normalized values of RMS, peak-to-peak amplitude, crest factor, kurtosis, energy, "
        "spectral energy, high-frequency ratio, and speed-normalized RMS. Positive deviations from the healthy baseline "
        "are averaged to produce a single score."
    )
    add_picture(document, "severity_score_by_file.png", "Figure 18. Estimated file-level severity score.", 6.5)
    severity_table = severity[["file", "health_name", "speed_name", "severity_score", "estimated_level"]].copy()
    add_table(document, severity_table, ["file", "health_name", "speed_name", "severity_score", "estimated_level"], "Table 7. Estimated relative severity by file.")
    document.add_paragraph(
        "The severity labels should be interpreted carefully. Good/normal, mild abnormal, and severe abnormal are "
        "relative vibration abnormality categories. They are not physical defect-size labels and should not be used as "
        "remaining useful life predictions."
    )

    document.add_heading("9. Limitations", level=1)
    add_bullets(
        document,
        [
            "The dataset contains repeated trials, not progressive degradation history.",
            "There are no true mild, moderate, or severe fault labels.",
            "The severity index is estimated from signal features and is not a physical damage measurement.",
            "The encoder speed unit should be confirmed before making strong physical interpretations of order values.",
            "File-level health results are strong, but the dataset is still small at the experiment-setting level.",
        ],
    )

    document.add_heading("10. Conclusion", level=1)
    document.add_paragraph(
        "The project successfully builds a bearing condition monitoring workflow for time-varying speed data. The raw "
        "signals were reduced into meaningful diagnostic features, and the classifier achieved perfect file-level health "
        "classification on the held-out trial-3 files. Signal plots and feature summaries support the interpretation that "
        "inner race faults are strongly separated in vibration features, while outer race faults require more careful "
        "frequency and envelope/order analysis."
    )
    document.add_paragraph(
        "The work should be described as fault diagnosis or condition monitoring, not full predictive maintenance with "
        "remaining useful life prediction. A relative severity index was included to rank abnormality, but a true "
        "mild/severe classifier would require labeled severity data."
    )

    document.add_heading("11. Appendix: Generated Files", level=1)
    appendix_table = pd.DataFrame(
        [
            ["analysis/", "MATLAB and Python scripts used to generate the complete workflow."],
            ["run_all_analysis.m", "One-command MATLAB runner for the analysis."],
            ["outputs/tables/window_features.csv", "Window-level extracted feature dataset."],
            ["outputs/tables/file_feature_summary.csv", "File-level feature summary."],
            ["outputs/tables/health_file_predictions.csv", "Trial-3 health classification predictions."],
            ["outputs/tables/file_severity_scores.csv", "Estimated file-level severity scores."],
            ["outputs/figures/", "All generated plots and confusion matrices."],
            ["outputs/report/Bearing_Analysis_Report.md", "Markdown version of the report."],
            ["outputs/report/Bearing_Fault_Diagnosis_Detailed_Report.docx", "This Word report."],
        ],
        columns=["path", "description"],
    )
    add_table(document, appendix_table, ["path", "description"], "Table 8. Main generated project files.")

    document.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = build_report()
    print(path)
